"""Word document parser using python-docx.

Extracts text, tables, and structure from .docx files.
"""

from pathlib import Path
from typing import Optional, List

from loguru import logger

from src.ingestion.parsers.base_parser import BaseParser
from src.schema import DocumentType, ParsedDocument


class WordParser(BaseParser):
    """Parser for Microsoft Word documents (.docx).
    
    Uses python-docx to extract:
    - Paragraphs with styles (headings, body text)
    - Tables with structure
    - Metadata (title, author, etc.)
    """
    
    SUPPORTED_EXTENSIONS = {".docx", ".doc"}
    
    def can_parse(self, file_path: str | Path) -> bool:
        """Check if file is a Word document."""
        return Path(file_path).suffix.lower() in self.SUPPORTED_EXTENSIONS
    
    def parse(
        self, file_path: str | Path, document_type: Optional[DocumentType] = None
    ) -> ParsedDocument:
        """Parse Word document.
        
        Args:
            file_path: Path to Word file
            document_type: Optional document type override
            
        Returns:
            ParsedDocument with structured text
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"Word file not found: {file_path}")
        
        logger.info(f"[WordParser] Parsing {file_path.name}...")
        
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")
        
        doc = Document(str(file_path))
        
        paragraphs = []
        headers = []
        tables_text = []
        
        # Extract paragraphs and identify headers by style
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
                
            paragraphs.append(text)
            
            # Check if heading style
            if para.style and "Heading" in para.style.name:
                headers.append(text)
        
        # Extract tables
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                table_rows.append(" | ".join(row_cells))
            
            if table_rows:
                tables_text.append("\n".join(table_rows))
                paragraphs.extend(table_rows)
        
        # Combine all text
        full_text = "\n\n".join(paragraphs)
        
        # Get metadata
        core_props = doc.core_properties
        metadata = {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else "",
        }
        
        # Detect document type if not provided
        if document_type is None:
            document_type = self._detect_document_type(full_text, file_path.name)
        
        logger.success(f"[WordParser] Parsed {file_path.name}: {len(full_text)} chars, {len(paragraphs)} paragraphs")
        
        return ParsedDocument(
            file_path=str(file_path),
            file_name=file_path.name,
            document_type=document_type,
            content=full_text,
            paragraphs=paragraphs,
            page_count=1,  # Word doesn't have reliable page count without rendering
            headers=headers,
            tables=tables_text,
            metadata=metadata,
        )
    
    def _detect_document_type(self, content: str, filename: str) -> DocumentType:
        """Detect document type from content and filename."""
        content_lower = content.lower()
        filename_lower = filename.lower()
        
        # Simple heuristics
        if any(word in content_lower for word in ["witness statement", "i hereby declare"]):
            return DocumentType.WITNESS_STATEMENT
        elif any(word in content_lower for word in ["court order", "judgment", "the court orders"]):
            return DocumentType.COURT_FILING
        elif any(word in content_lower for word in ["contract", "agreement", "parties agree"]):
            return DocumentType.CONTRACT
        elif "letter" in filename_lower or content_lower.startswith("dear "):
            return DocumentType.CORRESPONDENCE
        else:
            return DocumentType.GENERIC
